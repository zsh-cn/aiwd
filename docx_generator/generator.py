"""DOCX 生成模块"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocxGenerator:
    """DOCX 文件生成器"""

    ILLEGAL_CHARS_PATTERN = re.compile(r'[\\/:*?"<>|]')

    @classmethod
    def sanitize_filename(cls, name: str, max_len: int = 80) -> str:
        """清理非法文件名字符"""
        name = cls.ILLEGAL_CHARS_PATTERN.sub("_", name)
        name = name.strip().strip(".")
        if not name:
            name = "未命名文档"
        if len(name) > max_len:
            name = name[:max_len]
        return name

    @classmethod
    def generate(cls, title: str, content: str, output_dir: str) -> str:
        """生成单个 DOCX 文件，返回生成的文件路径"""
        os.makedirs(output_dir, exist_ok=True)

        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "宋体"
        font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(6)

        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.name = "黑体"

        paragraphs = content.split("\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0.74)
                run = p.add_run(para_text)
                run.font.name = "宋体"
                run.font.size = Pt(12)

        filename = cls.sanitize_filename(title) + ".docx"
        filepath = os.path.join(output_dir, filename)

        counter = 1
        base, ext = os.path.splitext(filename)
        while os.path.exists(filepath):
            filename = f"{base}_{counter}{ext}"
            filepath = os.path.join(output_dir, filename)
            counter += 1

        doc.save(filepath)
        return filepath