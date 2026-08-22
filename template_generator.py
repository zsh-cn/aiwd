import os
import tempfile


_PAPER_SIZES = {
    "A4": (21.0, 29.7),
    "A3": (29.7, 42.0),
    "A5": (14.8, 21.0),
    "LETTER": (21.59, 27.94),
    "LEGAL": (21.59, 35.56),
    "EXECUTIVE": (18.41, 26.67),
    "B5": (17.6, 25.7),
}


def generate_reference_doc(layout_config: dict, output_path: str = None) -> str:
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("请先安装 python-docx: pip install python-docx")

    if output_path is None:
        output_path = os.path.join(
            tempfile.gettempdir(), f"reference_doc_{abs(hash(str(layout_config)))}.docx"
        )

    doc = Document()
    styles_config = layout_config.get("styles", {})
    page_config = layout_config.get("page", {})
    footnotes_config = layout_config.get("footnotes", {})
    endnotes_config = layout_config.get("endnotes", {})

    _setup_page(doc, page_config, Cm, WD_ORIENT, qn)
    _configure_default_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH)
    _configure_title_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE)
    _configure_heading_styles(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE)
    _configure_body_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH)
    _configure_quote_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE)
    _configure_list_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE)
    _configure_code_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE, OxmlElement)
    _configure_inline_code_style(doc, styles_config, Pt, RGBColor, qn, WD_STYLE_TYPE, OxmlElement)
    _configure_hr_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE, OxmlElement)
    _configure_task_list_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE)
    _configure_definition_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE)
    _configure_table_style(doc, styles_config, Pt, Cm, RGBColor, qn, WD_STYLE_TYPE, OxmlElement)
    _configure_link_style(doc, styles_config, Pt, RGBColor, qn, WD_STYLE_TYPE)
    _configure_image_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE)
    _configure_footnotes_style(doc, footnotes_config, Pt, RGBColor, qn, WD_STYLE_TYPE, OxmlElement)
    _configure_endnotes_style(doc, endnotes_config, Pt, RGBColor, qn, WD_STYLE_TYPE, OxmlElement)
    _setup_headers_footers(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, OxmlElement)

    doc.save(output_path)
    return output_path


def _setup_page(doc, page_config: dict, Cm, WD_ORIENT, qn):
    for section in doc.sections:
        section.top_margin = Cm(page_config.get("margin_top", 2.54))
        section.bottom_margin = Cm(page_config.get("margin_bottom", 2.54))
        section.left_margin = Cm(page_config.get("margin_left", 3.18))
        section.right_margin = Cm(page_config.get("margin_right", 3.18))

        paper_size = page_config.get("paper_size", "A4").upper()
        if paper_size in _PAPER_SIZES:
            w, h = _PAPER_SIZES[paper_size]
        else:
            w, h = _PAPER_SIZES["A4"]

        orientation = page_config.get("orientation", "portrait").lower()
        if orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Cm(h)
            section.page_height = Cm(w)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Cm(w)
            section.page_height = Cm(h)

        columns = page_config.get("columns", 1)
        if columns and columns > 1:
            _apply_columns(section, columns, qn)

        page_border_color = page_config.get("page_border_color")
        if page_border_color:
            _apply_page_border(section, page_border_color,
                               page_config.get("page_border_width", 1.0), qn)

        background_color = page_config.get("background_color")
        if background_color:
            _apply_page_background(section, background_color, qn)


def _apply_columns(section, num_columns: int, qn):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = sectPr.makeelement(qn('w:cols'), {})
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num_columns))
    cols.set(qn('w:space'), '425')


def _apply_page_border(section, color: str, width_pt: float, qn):
    from lxml import etree
    sectPr = section._sectPr
    for existing in sectPr.findall(qn('w:pgBorders')):
        sectPr.remove(existing)
    pgBorders = etree.SubElement(sectPr, qn('w:pgBorders'))
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for edge in ['top', 'left', 'bottom', 'right']:
        border = etree.SubElement(pgBorders, qn(f'w:{edge}'))
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(int(width_pt * 8)))
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), color.lstrip('#'))


def _apply_page_background(section, color: str, qn):
    from lxml import etree
    sectPr = section._sectPr
    for existing in sectPr.findall(qn('w:background')):
        sectPr.remove(existing)
    bg = etree.SubElement(sectPr, qn('w:background'))
    bg.set(qn('w:color'), color.lstrip('#'))


def _set_font(run_or_style, Pt, RGBColor, qn,
              font_name=None, size=None, color=None,
              bold=None, italic=None, strike=None, underline=None,
              superscript=None, subscript=None, character_spacing=None,
              shadow=None, outline=None, emboss=None, imprint=None, glow=None):
    if font_name:
        run_or_style.font.name = font_name
        try:
            run_or_style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        except Exception:
            pass
    if size:
        run_or_style.font.size = Pt(size)
    if color:
        hex_color = color.lstrip('#')
        if len(hex_color) == 6:
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
            run_or_style.font.color.rgb = RGBColor(r, g, b)
    if bold is not None:
        run_or_style.font.bold = bold
    if italic is not None:
        run_or_style.font.italic = italic
    if strike is not None:
        run_or_style.font.strike = strike
    if underline is not None:
        run_or_style.font.underline = underline
    if superscript is not None:
        run_or_style.font.superscript = superscript
    if subscript is not None:
        run_or_style.font.subscript = subscript
    if character_spacing is not None and character_spacing != 0:
        _set_character_spacing(run_or_style, character_spacing, qn)
    rPr = run_or_style._element.get_or_add_rPr()
    if shadow is not None:
        _set_bool_prop(rPr, qn, 'w:shadow', shadow)
    if outline is not None:
        _set_bool_prop(rPr, qn, 'w:outline', outline)
    if emboss is not None:
        _set_bool_prop(rPr, qn, 'w:emboss', emboss)
    if imprint is not None:
        _set_bool_prop(rPr, qn, 'w:imprint', imprint)
    if glow is not None:
        _set_bool_prop(rPr, qn, 'w:glow', glow)


def _set_character_spacing(style, spacing_pt: float, qn):
    try:
        rPr = style._element.get_or_add_rPr()
        for existing in rPr.findall(qn('w:spacing')):
            rPr.remove(existing)
        from lxml import etree
        spacing_elem = etree.SubElement(rPr, qn('w:spacing'))
        spacing_elem.set(qn('w:charSpace'), str(int(spacing_pt * 20)))
    except Exception:
        pass


def _set_bool_prop(parent, qn, tag: str, value: bool):
    from lxml import etree
    for existing in parent.findall(qn(tag)):
        parent.remove(existing)
    if value:
        elem = etree.SubElement(parent, qn(tag))
        elem.set(qn('w:val'), '1')


def _set_paragraph_format(paragraph_format, config: dict, Pt, WD_ALIGN_PARAGRAPH):
    if "alignment" in config:
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        paragraph_format.alignment = align_map.get(config["alignment"], WD_ALIGN_PARAGRAPH.LEFT)

    if "space_before" in config:
        paragraph_format.space_before = Pt(config["space_before"])
    if "space_after" in config:
        paragraph_format.space_after = Pt(config["space_after"])
    if "line_spacing" in config:
        paragraph_format.line_spacing = config["line_spacing"]
    if "first_line_indent" in config:
        paragraph_format.first_line_indent = Pt(config["first_line_indent"])
    if "left_indent" in config:
        paragraph_format.left_indent = Pt(config["left_indent"])
    if "right_indent" in config:
        paragraph_format.right_indent = Pt(config["right_indent"])
    if "page_break_before" in config:
        paragraph_format.page_break_before = config["page_break_before"]
    if "page_break_after" in config:
        paragraph_format.page_break_after = config["page_break_after"]

    if "tab_stops" in config and isinstance(config["tab_stops"], list):
        _set_tab_stops(paragraph_format, config["tab_stops"], Pt, WD_ALIGN_PARAGRAPH)


def _set_tab_stops(paragraph_format, tab_stops: list, Pt, WD_ALIGN_PARAGRAPH):
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    for ts in tab_stops:
        if isinstance(ts, dict) and "position" in ts:
            position = Pt(ts["position"])
            alignment = align_map.get(ts.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
            paragraph_format.tab_stops.add_tab_stop(position, alignment)


def _get_or_create_style(doc, name, style_type):
    try:
        return doc.styles[name]
    except KeyError:
        pass
    try:
        return doc.styles.add_style(name, style_type)
    except Exception:
        return doc.styles[name]


def _apply_font_config(style, config: dict, Pt, RGBColor, qn):
    font_name = config.get("font", config.get("name"))
    _set_font(
        style, Pt, RGBColor, qn,
        font_name=font_name,
        size=config.get("size"),
        color=config.get("color"),
        bold=config.get("bold"),
        italic=config.get("italic"),
        strike=config.get("strike"),
        underline=config.get("underline"),
        superscript=config.get("superscript"),
        subscript=config.get("subscript"),
        character_spacing=config.get("character_spacing"),
        shadow=config.get("shadow"),
        outline=config.get("outline"),
        emboss=config.get("emboss"),
        imprint=config.get("imprint"),
        glow=config.get("glow"),
    )


def _apply_para_config(style, config: dict, Pt, WD_ALIGN_PARAGRAPH):
    _set_paragraph_format(style.paragraph_format, config, Pt, WD_ALIGN_PARAGRAPH)


def _configure_default_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH):
    default_config = styles_config.get("default_font", {})
    body_config = styles_config.get("body", {})

    style = doc.styles['Normal']
    _apply_font_config(style, default_config, Pt, RGBColor, qn)

    if body_config:
        _apply_para_config(style, body_config, Pt, WD_ALIGN_PARAGRAPH)


def _configure_title_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE):
    title_config = styles_config.get("title", {})
    if not title_config:
        return

    style = _get_or_create_style(doc, 'Title', WD_STYLE_TYPE.PARAGRAPH)
    _apply_font_config(style, title_config, Pt, RGBColor, qn)
    _apply_para_config(style, title_config, Pt, WD_ALIGN_PARAGRAPH)


def _configure_heading_styles(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE):
    heading_map = {
        'h1': 'Heading 1',
        'h2': 'Heading 2',
        'h3': 'Heading 3',
        'h4': 'Heading 4',
        'h5': 'Heading 5',
        'h6': 'Heading 6',
    }

    for config_key, style_name in heading_map.items():
        config = styles_config.get(config_key, {})
        if not config:
            continue

        style = _get_or_create_style(doc, style_name, WD_STYLE_TYPE.PARAGRAPH)
        _apply_font_config(style, config, Pt, RGBColor, qn)
        _apply_para_config(style, config, Pt, WD_ALIGN_PARAGRAPH)


def _configure_body_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH):
    body_config = styles_config.get("body", {})
    if not body_config:
        return

    style = doc.styles['Normal']
    _apply_font_config(style, body_config, Pt, RGBColor, qn)
    _apply_para_config(style, body_config, Pt, WD_ALIGN_PARAGRAPH)


def _configure_quote_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE):
    quote_config = styles_config.get("quote", {})
    if not quote_config:
        return

    style = _get_or_create_style(doc, 'Quote', WD_STYLE_TYPE.PARAGRAPH)
    _apply_font_config(style, quote_config, Pt, RGBColor, qn)
    _apply_para_config(style, quote_config, Pt, WD_ALIGN_PARAGRAPH)


def _configure_list_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE):
    base_list_config = styles_config.get("list", {})
    nested_config = styles_config.get("nested_list", {})

    if base_list_config:
        for style_name in ['List Bullet', 'List Number']:
            style = _get_or_create_style(doc, style_name, WD_STYLE_TYPE.PARAGRAPH)
            _apply_font_config(style, base_list_config, Pt, RGBColor, qn)
            _apply_para_config(style, base_list_config, Pt, WD_ALIGN_PARAGRAPH)

    if nested_config:
        for level in [2, 3]:
            for prefix in ['List Bullet', 'List Number']:
                style_name = f'{prefix} {level}'
                style = _get_or_create_style(doc, style_name, WD_STYLE_TYPE.PARAGRAPH)
                level_config = nested_config.get(f'level{level}',
                                                 nested_config.get(f'level_{level}',
                                                                   nested_config))
                _apply_font_config(style, level_config, Pt, RGBColor, qn)
                _apply_para_config(style, level_config, Pt, WD_ALIGN_PARAGRAPH)


def _configure_code_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE, OxmlElement):
    code_config = styles_config.get("code", {})
    if not code_config:
        return

    style = _get_or_create_style(doc, 'Code', WD_STYLE_TYPE.PARAGRAPH)
    _apply_font_config(style, code_config, Pt, RGBColor, qn)
    _apply_para_config(style, code_config, Pt, WD_ALIGN_PARAGRAPH)

    bg_color = code_config.get("background_color")
    if bg_color:
        _set_style_shading(style, bg_color, qn, OxmlElement)


def _configure_inline_code_style(doc, styles_config, Pt, RGBColor, qn, WD_STYLE_TYPE, OxmlElement):
    inline_code_config = styles_config.get("inline_code", {})
    if not inline_code_config:
        return

    style = _get_or_create_style(doc, 'Source Code', WD_STYLE_TYPE.PARAGRAPH)
    _apply_font_config(style, inline_code_config, Pt, RGBColor, qn)

    bg_color = inline_code_config.get("background_color")
    if bg_color:
        _set_style_shading(style, bg_color, qn, OxmlElement)


def _configure_hr_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE, OxmlElement):
    hr_config = styles_config.get("hr", {})
    if not hr_config:
        return

    style = _get_or_create_style(doc, 'Horizontal Rule', WD_STYLE_TYPE.PARAGRAPH)

    space_before = hr_config.get("space_before", 12)
    space_after = hr_config.get("space_after", 12)
    style.paragraph_format.space_before = Pt(space_before)
    style.paragraph_format.space_after = Pt(space_after)

    color = hr_config.get("color", "#999999")
    width = hr_config.get("width", 1.0)
    _set_paragraph_bottom_border(style, color, width, qn, OxmlElement)


def _configure_task_list_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE):
    task_config = styles_config.get("task_list", {})
    if not task_config:
        return

    style = _get_or_create_style(doc, 'List Bullet', WD_STYLE_TYPE.PARAGRAPH)
    _apply_font_config(style, task_config, Pt, RGBColor, qn)
    _apply_para_config(style, task_config, Pt, WD_ALIGN_PARAGRAPH)


def _configure_definition_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE):
    def_config = styles_config.get("definition", {})
    if not def_config:
        return

    term_font = def_config.get("term_font", "黑体")
    term_size = def_config.get("term_size", 11)
    term_color = def_config.get("term_color", "#1a202c")
    term_bold = def_config.get("term_bold", True)

    def_font = def_config.get("definition_font", "宋体")
    def_size = def_config.get("definition_size", 11)
    def_color = def_config.get("definition_color", "#000000")
    def_italic = def_config.get("definition_italic", False)
    left_indent = def_config.get("left_indent", 44)

    term_style = _get_or_create_style(doc, 'Definition Term', WD_STYLE_TYPE.PARAGRAPH)
    _set_font(term_style, Pt, RGBColor, qn,
              font_name=term_font, size=term_size, color=term_color,
              bold=term_bold)
    term_style.paragraph_format.space_before = Pt(6)
    term_style.paragraph_format.space_after = Pt(3)

    def_style = _get_or_create_style(doc, 'Definition', WD_STYLE_TYPE.PARAGRAPH)
    _set_font(def_style, Pt, RGBColor, qn,
              font_name=def_font, size=def_size, color=def_color,
              italic=def_italic)
    def_style.paragraph_format.left_indent = Pt(left_indent)
    def_style.paragraph_format.space_before = Pt(3)
    def_style.paragraph_format.space_after = Pt(3)


def _configure_table_style(doc, styles_config, Pt, Cm, RGBColor, qn, WD_STYLE_TYPE, OxmlElement):
    table_config = styles_config.get("table", {})
    if not table_config:
        return

    style = _get_or_create_style(doc, 'Table Grid', WD_STYLE_TYPE.TABLE)
    _apply_font_config(style, table_config, Pt, RGBColor, qn)

    if table_config.get("border", True):
        _set_table_borders(
            style,
            color=table_config.get("border_color", "#000000"),
            width=table_config.get("border_width", 0.5),
            qn=qn,
            OxmlElement=OxmlElement,
        )
    else:
        _remove_table_borders(style, qn, OxmlElement)

    row_height = table_config.get("row_height")
    if row_height:
        _set_table_row_height(style, row_height, Cm, qn, OxmlElement)

    has_cell_margins = any(
        table_config.get(k) is not None
        for k in ["cell_margin_top", "cell_margin_bottom", "cell_margin_left", "cell_margin_right"]
    )
    if has_cell_margins:
        _set_table_cell_margins(style, table_config, Cm, qn, OxmlElement)

    cell_shading = table_config.get("cell_shading")
    if cell_shading:
        _set_style_shading(style, cell_shading, qn, OxmlElement)

    header_row = table_config.get("header_row", {})
    if header_row:
        _set_table_header_row(style, header_row, Pt, RGBColor, qn, OxmlElement)


def _configure_link_style(doc, styles_config, Pt, RGBColor, qn, WD_STYLE_TYPE):
    link_config = styles_config.get("link", {})
    if not link_config:
        return

    style = _get_or_create_style(doc, 'Hyperlink', WD_STYLE_TYPE.PARAGRAPH)
    _apply_font_config(style, link_config, Pt, RGBColor, qn)


def _configure_image_style(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE):
    image_config = styles_config.get("image", {})
    if not image_config:
        return

    style = _get_or_create_style(doc, 'Image', WD_STYLE_TYPE.PARAGRAPH)

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    style.paragraph_format.alignment = align_map.get(
        image_config.get("alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER
    )

    caption_font = image_config.get("caption_font", "宋体")
    caption_size = image_config.get("caption_size", 9)
    caption_color = image_config.get("caption_color", "#666666")

    _set_font(style, Pt, RGBColor, qn,
              font_name=caption_font, size=caption_size, color=caption_color)


def _configure_footnotes_style(doc, config, Pt, RGBColor, qn, WD_STYLE_TYPE, OxmlElement):
    if not config:
        return
    try:
        style = doc.styles['Footnote Text']
    except KeyError:
        return
    _apply_font_config(style, config, Pt, RGBColor, qn)


def _configure_endnotes_style(doc, config, Pt, RGBColor, qn, WD_STYLE_TYPE, OxmlElement):
    if not config:
        return
    try:
        style = doc.styles['Endnote Text']
    except KeyError:
        return
    _apply_font_config(style, config, Pt, RGBColor, qn)


def _setup_headers_footers(doc, styles_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, OxmlElement):
    header_config = styles_config.get("header", {})
    footer_config = styles_config.get("footer", {})

    for section in doc.sections:
        if header_config and header_config.get("text"):
            _setup_header(section, header_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH)

        if footer_config and (footer_config.get("text") or footer_config.get("page_number")):
            _setup_footer(section, footer_config, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, OxmlElement)


def _setup_header(section, config: dict, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH):
    header = section.header
    if header.paragraphs:
        para = header.paragraphs[0]
    else:
        para = header.add_paragraph()

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    para.alignment = align_map.get(config.get("alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER)

    run = para.add_run(config.get("text", ""))
    font_name = config.get("font", "宋体")
    font_size = config.get("size", 9)
    font_color = config.get("color", "#666666")

    run.font.name = font_name
    try:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass
    run.font.size = Pt(font_size)
    hex_color = font_color.lstrip('#')
    if len(hex_color) == 6:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def _setup_footer(section, config: dict, Pt, RGBColor, qn, WD_ALIGN_PARAGRAPH, OxmlElement):
    footer = section.footer
    if footer.paragraphs:
        para = footer.paragraphs[0]
    else:
        para = footer.add_paragraph()

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    para.alignment = align_map.get(config.get("alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER)

    font_name = config.get("font", "宋体")
    font_size = config.get("size", 9)
    font_color = config.get("color", "#666666")

    def _add_footer_run(text=""):
        run = para.add_run(text)
        run.font.name = font_name
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        except Exception:
            pass
        run.font.size = Pt(font_size)
        hex_c = font_color.lstrip('#')
        if len(hex_c) == 6:
            r = int(hex_c[0:2], 16)
            g = int(hex_c[2:4], 16)
            b = int(hex_c[4:6], 16)
            run.font.color.rgb = RGBColor(r, g, b)
        return run

    footer_text = config.get("text", "")
    page_number = config.get("page_number", False)

    if footer_text:
        _add_footer_run(footer_text)

    if page_number:
        if footer_text:
            _add_footer_run("  ")
        _add_page_number_field(para, font_name, font_size, font_color, Pt, RGBColor, qn, OxmlElement)


def _add_page_number_field(para, font_name, font_size, font_color, Pt, RGBColor, qn, OxmlElement):
    run1 = para.add_run()
    run1.font.name = font_name
    try:
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass
    run1.font.size = Pt(font_size)
    hex_c = font_color.lstrip('#')
    if len(hex_c) == 6:
        r = int(hex_c[0:2], 16)
        g = int(hex_c[2:4], 16)
        b = int(hex_c[4:6], 16)
        run1.font.color.rgb = RGBColor(r, g, b)

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run1._element.append(fldChar_begin)

    run2 = para.add_run()
    run2.font.name = font_name
    try:
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass
    run2.font.size = Pt(font_size)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run2._element.append(instrText)

    run3 = para.add_run()
    run3.font.name = font_name
    try:
        run3._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass
    run3.font.size = Pt(font_size)

    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    run3._element.append(fldChar_sep)

    run4 = para.add_run()
    run4.font.name = font_name
    try:
        run4._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass
    run4.font.size = Pt(font_size)
    run4.text = "1"

    run5 = para.add_run()
    run5.font.name = font_name
    try:
        run5._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass
    run5.font.size = Pt(font_size)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run5._element.append(fldChar_end)


def _set_table_borders(style, color: str, width: float, qn, OxmlElement):
    tbl = style.element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)

    tblBorders = OxmlElement('w:tblBorders')
    border_sz = str(int(width * 8))
    border_color = color.lstrip('#')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), border_sz)
        border.set(qn('w:color'), border_color)
        tblBorders.append(border)

    tblPr.append(tblBorders)


def _remove_table_borders(style, qn, OxmlElement):
    tbl = style.element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)

    tblPr.append(tblBorders)


def _set_table_row_height(style, height_cm: float, Cm, qn, OxmlElement):
    tbl = style.element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    trPr = OxmlElement('w:trPr')
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

    for existing in tblPr.findall(qn('w:tblBorders')):
        pass
    tblPr.append(trPr)


def _set_table_cell_margins(style, table_config: dict, Cm, qn, OxmlElement):
    tbl = style.element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    for existing in tblPr.findall(qn('w:tblCellMar')):
        tblPr.remove(existing)

    tblCellMar = OxmlElement('w:tblCellMar')
    for side, key in [('top', 'cell_margin_top'), ('bottom', 'cell_margin_bottom'),
                       ('left', 'cell_margin_left'), ('right', 'cell_margin_right')]:
        val = table_config.get(key)
        if val is not None:
            elem = OxmlElement(f'w:{side}')
            elem.set(qn('w:w'), str(int(val * 567)))
            elem.set(qn('w:type'), 'dxa')
            tblCellMar.append(elem)

    tblPr.append(tblCellMar)


def _set_style_shading(style, color: str, qn, OxmlElement):
    try:
        pPr = style.element.get_or_add_pPr()
        for existing in pPr.findall(qn('w:shd')):
            pPr.remove(existing)

        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color.lstrip('#'))
        pPr.append(shd)
    except Exception:
        pass


def _set_table_cell_shading(style, color: str, qn, OxmlElement):
    try:
        tbl = style.element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        for existing in tblPr.findall(qn('w:tblCellSpacing')):
            tblPr.remove(existing)

        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color.lstrip('#'))

        try:
            from lxml import etree
            condFormatting = etree.SubElement(tblPr, qn('w:tblCondFmt'))
            condFormatting.set(qn('w:type'), 'wholeTable')
            condFormatting.append(shd)
        except Exception:
            pass
    except Exception:
        pass


def _set_paragraph_bottom_border(style, color: str, width_pt: float, qn, OxmlElement):
    try:
        pPr = style.element.get_or_add_pPr()
        for existing in pPr.findall(qn('w:pBdr')):
            pPr.remove(existing)

        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(int(width_pt * 8)))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color.lstrip('#'))
        pBdr.append(bottom)
        pPr.append(pBdr)
    except Exception:
        pass


def _set_table_header_row(style, header_config: dict, Pt, RGBColor, qn, OxmlElement):
    try:
        tbl = style.element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        tblLook = OxmlElement('w:tblLook')
        tblLook.set(qn('w:val'), '04A0')
        tblLook.set(qn('w:firstRow'), '1')
        tblLook.set(qn('w:lastRow'), '0')
        tblLook.set(qn('w:firstColumn'), '0')
        tblLook.set(qn('w:lastColumn'), '0')
        tblLook.set(qn('w:noHBand'), '0')
        tblLook.set(qn('w:noVBand'), '1')
        for existing in tblPr.findall(qn('w:tblLook')):
            tblPr.remove(existing)
        tblPr.append(tblLook)

        header_bg = header_config.get("background_color")
        if header_bg:
            try:
                from lxml import etree
                ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                condFormatting = etree.SubElement(tblPr, qn('w:tblCondFmt'))
                condFormatting.set(qn('w:type'), 'firstRow')
                shd = etree.SubElement(condFormatting, qn('w:shd'))
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), header_bg.lstrip('#'))
            except Exception:
                pass
    except Exception:
        pass